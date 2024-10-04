import io
import PyPDF2
import docx

def handle_uploaded_file(uploaded_file):
    if uploaded_file.type == "text/plain":
        content = uploaded_file.read().decode('utf-8')
    elif uploaded_file.type == "application/pdf":
        content = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        # Leggi il documento Word
        doc = docx.Document(uploaded_file)
        full_text = []

        # Estrai il testo dai paragrafi
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Estrai il testo dalle tabelle
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                full_text.append(' | '.join(row_data))

        content = '\n'.join(full_text)
        return content
    else:
        content = ""

    return content

def extract_text_from_pdf(uploaded_file):
    pdf_reader = PyPDF2.PdfReader(uploaded_file)
    text = ""
    for page_num in range(len(pdf_reader.pages)):
        text += pdf_reader.pages[page_num].extract_text()
    return text

def extract_text_from_docx(uploaded_file):
    doc = docx.Document(uploaded_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text
