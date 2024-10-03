import streamlit as st
from agents.agent1 import restructure_aru
from agents.agent2 import calculate_ufp
from agents.supervisor_agent import supervise_process
import logging
from dotenv import load_dotenv
import os
import PyPDF2
from io import BytesIO
from docx import Document
import base64
import markdown2
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.units import inch
import re  # Importa il modulo re per le espressioni regolari

# Carica le variabili d'ambiente
load_dotenv()

# Funzione per resettare il file di log
def reset_log():
    open('app.log', 'w').close()

# Configura il logging
def configure_logging():
    logging.basicConfig(filename='app.log', level=logging.INFO, format='%(asctime)s %(message)s')

# Funzione per leggere il contenuto di un file PDF
def read_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

# Funzione per leggere il contenuto di un file DOCX
def read_docx(file):
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

# Funzione per gestire lo streaming delle risposte e visualizzarle come markdown
def display_markdown_stream(generator, placeholder, initial_text=""):
    full_text = initial_text
    for chunk in generator:
        full_text += chunk
        placeholder.markdown(full_text, unsafe_allow_html=True)
    return full_text

# Funzione per creare un file DOCX dal testo markdown
def create_docx_from_markdown(markdown_text, file_name="document.docx"):
    doc = Document()
    html = markdown2.markdown(markdown_text)
    soup = BeautifulSoup(html, 'html.parser')
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol']):
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Number')
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Funzione per creare un file PDF con ReportLab
def create_pdf_from_markdown(markdown_text, file_name="document.pdf"):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    story = []

    # Converti markdown in HTML, quindi in PDF con ReportLab
    html = markdown2.markdown(markdown_text)
    soup = BeautifulSoup(html, 'html.parser')

    # Crea uno stile di base per il PDF
    styles = getSampleStyleSheet()

    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol']):
        if element.name == 'h1':
            story.append(Paragraph(f"<b>{element.text}</b>", styles['Heading1']))
        elif element.name == 'h2':
            story.append(Paragraph(f"<b>{element.text}</b>", styles['Heading2']))
        elif element.name == 'h3':
            story.append(Paragraph(f"<b>{element.text}</b>", styles['Heading3']))
        elif element.name == 'p':
            story.append(Paragraph(element.text, styles['BodyText']))
        elif element.name == 'ul':
            for li in element.find_all('li'):
                story.append(Paragraph(f"- {li.text}", styles['BodyText']))
        elif element.name == 'ol':
            for idx, li in enumerate(element.find_all('li'), start=1):
                story.append(Paragraph(f"{idx}. {li.text}", styles['BodyText']))

    doc.build(story)
    buffer.seek(0)
    return buffer

def main():
    # Reset del file di log
    reset_log()

    # Configura il logging
    configure_logging()

    st.set_page_config(page_title="Calcolo Function Point", layout="wide")
    st.title("Calcolo dei Function Point da ARU")

    # Inizializza le variabili di sessione
    if 'aru_content' not in st.session_state:
        st.session_state['aru_content'] = ''
    if 'aru_ristrutturato' not in st.session_state:
        st.session_state['aru_ristrutturato'] = ''
    if 'ufp_results' not in st.session_state:
        st.session_state['ufp_results'] = ''
    if 'supervision_report' not in st.session_state:
        st.session_state['supervision_report'] = ''
    if 'total_ufp' not in st.session_state:
        st.session_state['total_ufp'] = None

    # Step 1: Caricamento ARU
    st.header("1. Carica il tuo documento ARU")
    uploaded_file = st.file_uploader("Scegli un file", type=["pdf", "doc", "docx", "txt"])

    if uploaded_file is not None:
        # Estrai il testo dal file caricato in base al formato
        if uploaded_file.type == "application/pdf":
            aru_content = read_pdf(uploaded_file)
        elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            # Salva il file caricato temporaneamente
            with open("temp_uploaded_file.docx", "wb") as temp_file:
                temp_file.write(uploaded_file.getbuffer())
            # Leggi il contenuto del file DOCX
            aru_content = read_docx("temp_uploaded_file.docx")
            # Rimuovi il file temporaneo
            os.remove("temp_uploaded_file.docx")
        else:  # Tratta il file come un semplice testo
            aru_content = uploaded_file.read().decode('utf-8', errors='ignore')

        st.success("Documento caricato con successo!")
        logging.info("Documento ARU caricato dall'utente.")

        # Salva il contenuto originale dell'ARU nella sessione
        st.session_state['aru_content'] = aru_content

    # Visualizza il contenuto dell'ARU se disponibile
    if st.session_state['aru_content']:
        st.subheader("Contenuto del documento ARU")
        st.text_area("Visualizza l'ARU originale", st.session_state['aru_content'], height=300)

        # Step 2: Revisione ARU
        st.header("2. Revisione dell'ARU")
        if st.button("Esegui Revisione"):
            st.subheader("ARU Revisionato")
            aru_area = st.empty()
            with st.spinner('Revisione in corso...'):
                aru_text = ''.join([chunk for chunk in restructure_aru(st.session_state['aru_content'])])
                st.session_state['aru_ristrutturato'] = aru_text
                logging.info("Revisione dell'ARU completata.")

                # Log della chiamata e della risposta dell'agente
                logging.info("Chiamata all'agente 1 (restructure_aru) con richiesta:")
                logging.info(st.session_state['aru_content'])
                logging.info("Risposta dall'agente 1:")
                logging.info(st.session_state['aru_ristrutturato'])

        # Visualizza l'ARU revisionato se disponibile
        if st.session_state.get('aru_ristrutturato'):
            st.subheader("ARU Revisionato")
            with st.expander("Visualizza l'ARU Revisionato"):
                st.markdown(st.session_state['aru_ristrutturato'], unsafe_allow_html=True)

            # Crea un file DOCX dal documento revisionato
            docx_buffer = create_docx_from_markdown(st.session_state['aru_ristrutturato'])

            # Pulsante per scaricare il file DOCX
            st.download_button(
                label="Scarica ARU Revisionato in formato DOCX",
                data=docx_buffer,
                file_name="ARU_revisionato.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # Crea un file PDF dal documento revisionato
            pdf_buffer = create_pdf_from_markdown(st.session_state['aru_ristrutturato'])

            # Pulsante per scaricare il file PDF
            st.download_button(
                label="Scarica ARU Revisionato in formato PDF",
                data=pdf_buffer,
                file_name="ARU_revisionato.pdf",
                mime="application/pdf"
            )

            # Step 3: Calcolo UFP
            st.header("3. Calcolo degli Unadjusted Function Points (UFP)")
            if st.button("Calcola UFP"):
                st.subheader("Risultati del Calcolo UFP")
                ufp_area = st.empty()
                with st.spinner('Calcolo UFP in corso...'):
                    ufp_text = ''.join([chunk for chunk in calculate_ufp(st.session_state['aru_ristrutturato'])])
                    st.session_state['ufp_results'] = ufp_text
                    logging.info("Calcolo UFP completato.")

                    # Log della chiamata e della risposta dell'agente
                    logging.info("Chiamata all'agente 2 (calculate_ufp) con richiesta:")
                    logging.info(st.session_state['aru_ristrutturato'])
                    logging.info("Risposta dall'agente 2:")
                    logging.info(st.session_state['ufp_results'])

                    # Estrai e salva il totale UFP utilizzando una regex
                    total_ufp_match = re.search(r"Totale UFP:\s*(\d+)", st.session_state['ufp_results'])
                    if total_ufp_match:
                        st.session_state['total_ufp'] = int(total_ufp_match.group(1))
                    else:
                        st.session_state['total_ufp'] = None
                        st.warning("Impossibile estrarre il totale UFP dai risultati.")

            # Visualizza i risultati UFP se disponibili
            if st.session_state['ufp_results']:
                st.subheader("Risultati del Calcolo UFP")
                with st.expander("Visualizza i Risultati del Calcolo UFP"):
                    st.markdown(st.session_state['ufp_results'], unsafe_allow_html=True)

                # Mostra il totale UFP
                st.subheader("Totale Unadjusted Function Points (UFP)")
                if st.session_state['total_ufp'] is not None:
                    st.success(f"**Totale UFP:** {st.session_state['total_ufp']}")
                else:
                    st.warning("Impossibile estrarre il totale UFP dai risultati.")

                # Step 4: Supervisione
                st.header("4. Supervisione del Processo")
                if st.button("Esegui Supervisione"):
                    st.subheader("Rapporto del Supervisore")
                    supervision_area = st.empty()
                    with st.spinner('Supervisione in corso...'):
                        supervision_text = ''.join([chunk for chunk in supervise_process(
                            st.session_state['aru_ristrutturato'],
                            st.session_state['ufp_results'],
                            st.session_state['total_ufp']
                        )])
                        st.session_state['supervision_report'] = supervision_text
                        logging.info("Supervisione completata.")

                        # Log della chiamata e della risposta del supervisore
                        logging.info("Chiamata al supervisore con richiesta:")
                        logging.info(f"ARU Revisionato: {st.session_state['aru_ristrutturato']}")
                        logging.info(f"Risultati UFP: {st.session_state['ufp_results']}")
                        logging.info(f"Totale UFP: {st.session_state['total_ufp']}")
                        logging.info("Risposta dal supervisore:")
                        logging.info(st.session_state['supervision_report'])

            # Visualizza il rapporto di supervisione se disponibile
            if st.session_state['supervision_report']:
                st.subheader("Rapporto del Supervisore")
                with st.expander("Visualizza il Rapporto del Supervisore"):
                    st.markdown(st.session_state['supervision_report'], unsafe_allow_html=True)

                # Crea un file DOCX dal rapporto di supervisione
                report_docx_buffer = create_docx_from_markdown(st.session_state['supervision_report'], file_name="Rapporto_Supervisore.docx")

                # Pulsante per scaricare il rapporto di supervisione in formato DOCX
                st.download_button(
                    label="Scarica Rapporto del Supervisore in formato DOCX",
                    data=report_docx_buffer,
                    file_name="Rapporto_Supervisore.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                # Crea un file PDF dal rapporto di supervisione
                report_pdf_buffer = create_pdf_from_markdown(st.session_state['supervision_report'])

                # Pulsante per scaricare il rapporto di supervisione in formato PDF
                st.download_button(
                    label="Scarica Rapporto del Supervisore in formato PDF",
                    data=report_pdf_buffer,
                    file_name="Rapporto_Supervisore.pdf",
                    mime="application/pdf"
                )

        # Mostra log di tutti i colloqui tra agenti
        st.header("Log agenti")
        with st.expander("Visualizza i Log degli agenti"):
            try:
                with open('app.log', 'r') as f:
                    logs = f.read()
                    st.text_area("", logs, height=300, key="log")
            except FileNotFoundError:
                st.warning("Nessun log disponibile al momento.")

if __name__ == "__main__":
    main()
